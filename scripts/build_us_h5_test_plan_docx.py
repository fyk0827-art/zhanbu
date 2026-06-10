from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = "/Users/administrator/Desktop/work/movie_work/占卜/问答平台/docs/美国市场H5测试计划_转化导向版.docx"


AD_CASES = [
    ("AD-01", "验证主卖点", "对比 Love clarity、Relationship advice、Spiritual guidance 三种广告切入点。", "CTR、CPC、LPV", "至少 1 个切入点 CTR 明显领先。"),
    ("AD-02", "验证素材形式", "对比真人口播、结果展示、H5 录屏三类素材。", "CTR、LPV、CPI", "找到最低获客成本素材形式。"),
    ("AD-03", "验证首屏钩子", "测试“Will he come back?”、“Get clarity now”、“Your personalized reading” 等前 3 秒钩子。", "CTR、3 秒播放率、LPV", "至少 1 个钩子点击效率显著领先。"),
    ("AD-04", "验证平台差异", "在 Meta、TikTok、Google 三平台测试同类素材。", "CPM、CTR、开始率、CAC", "筛出主力买量渠道。"),
    ("AD-05", "验证人群方向", "测试情感关系、自我探索、灵性兴趣、泛女性兴趣四类受众包。", "CTR、开始率、付费率", "锁定 1 到 2 个高质量人群包。"),
    ("AD-06", "验证价格前置", "广告中带“$0.99 first reading”与不带价格版本对比。", "CTR、LP 转化率、付费率", "判断价格前置是否改善流量质量。"),
]

PRODUCT_CASES = [
    ("PD-01", "验证首页定位", "首页分别强调 Tarot reading、Relationship clarity、AI guidance。", "首屏停留、开始率", "开始率最高版本胜出。"),
    ("PD-02", "验证转化路径", "先答题再出结果，与先输入问题再出结果两条路径对比。", "开始率、完成率、付费率", "选出完整付费漏斗转化最高路径。"),
    ("PD-03", "验证结果解锁方式", "直接展示部分结果，与摘要后引导付费解锁完整结果对比。", "到支付页转化率、付费率", "找到最强付费推进方式。"),
    ("PD-04", "验证首单价格", "测试 $0.99、$2.99、$4.99 三档首单价格。", "付费率、ARPPU、CAC", "确认最佳转化价格甜点。"),
    ("PD-05", "验证支付页信任组件", "比较加入用户评价、隐私承诺、价格说明前后的转化差异。", "支付页停留、付费率", "信任组件带来正向提升。"),
    ("PD-06", "验证支付透明度", "支付页清晰展示价格、续费说明、权益说明，与弱展示版本对比。", "付费率、退款率", "优先保留高转化低争议方案。"),
    ("PD-07", "验证注册节点", "先注册后体验，与游客先体验后支付前注册两种策略对比。", "注册率、完成率、付费率", "找出阻力最小方案。"),
    ("PD-08", "验证结果风格", "比较偏预测型文案与偏陪伴建议型文案对支付转化的影响。", "完成率、付费率", "明确最适合 H5 成交的话术。"),
]

USER_CASES = [
    ("US-01", "验证点击动机", "分析用户是被情感问题、未来结果还是低价试用吸引进来。", "CTR、开始率", "明确最强点击驱动力。"),
    ("US-02", "验证中途流失点", "定位用户在首页、答题页、结果页、支付页的主要流失节点。", "各节点流失率", "锁定首要转化断点。"),
    ("US-03", "验证支付障碍", "访谈未付费用户，梳理价格、信任、理解成本三类阻碍。", "支付阻碍分布、付费率", "明确最大付费阻碍。"),
    ("US-04", "验证价格敏感度", "观察不同价格带下用户的点击后行为与支付接受度。", "付费率、ARPPU", "确定最优价格区间。"),
    ("US-05", "验证信任障碍", "识别用户在哪一步开始怀疑结果真实性或支付安全性。", "节点流失率、访谈反馈", "找到主要信任问题。"),
    ("US-06", "验证内容偏好", "对比 Love、Career、Future、Self-growth 四类入口的成交效率。", "完成率、付费率", "选出高转化主题。"),
]

DATA_CASES = [
    ("DA-01", "验证埋点完整性", "校验广告点击到支付完成的 H5 全链路埋点。", "关键事件完整率", "关键事件完整率高于 95%。"),
    ("DA-02", "验证归因准确性", "核对 Meta、TikTok、Google 的点击、到站、支付归因口径。", "归因一致性", "误差控制在可接受范围。"),
    ("DA-03", "验证转化漏斗看板", "搭建 LPV、开始、完成、到支付页、支付成功的日常监控看板。", "漏斗可用性", "可每日定位主要流失节点。"),
    ("CM-01", "验证支付与文案合规", "检查是否出现绝对化承诺、误导性价格说明或敏感替代表述。", "审核通过率、退款率、投诉率", "素材和支付表达基础合规。"),
]

PRIORITY_CASES = [
    "AD-01 主卖点测试",
    "AD-02 素材形式测试",
    "AD-04 平台差异测试",
    "AD-05 人群方向测试",
    "PD-02 转化路径测试",
    "PD-03 结果解锁方式测试",
    "PD-04 首单价格测试",
    "PD-05 支付页信任组件测试",
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def format_run(run, size, bold=False, color="1F1F1F", font_name="Calibri"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def format_paragraph(paragraph, after=6, before=0, line=1.1, align=WD_ALIGN_PARAGRAPH.LEFT):
    pf = paragraph.paragraph_format
    pf.space_after = Pt(after)
    pf.space_before = Pt(before)
    pf.line_spacing = line
    paragraph.alignment = align


def add_title(doc, text, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    format_run(run, 22, bold=True, color="17365D")
    format_paragraph(p, after=8, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(subtitle)
    format_run(run2, 10, color="666666")
    format_paragraph(p2, after=18, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        format_run(run, 15, bold=True, color="2E74B5")
        format_paragraph(p, after=8, before=14)
    else:
        format_run(run, 12.5, bold=True, color="2E74B5")
        format_paragraph(p, after=6, before=10)


def add_body(doc, text, bullet=False):
    p = doc.add_paragraph(style="List Bullet" if bullet else None)
    run = p.add_run(text)
    format_run(run, 10.5)
    format_paragraph(p, after=6, line=1.15)


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    format_run(run, 10.5)
    format_paragraph(p, after=6, line=1.15)


def add_kv_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    widths = [Inches(1.55), Inches(4.95)]
    for key, value in rows:
        cells = table.add_row().cells
        set_row_cant_split(table.rows[-1])
        for idx, width in enumerate(widths):
            cells[idx].width = width
        set_cell_shading(cells[0], "F2F4F7")
        for cell in cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
        p1 = cells[0].paragraphs[0]
        p1.clear()
        r1 = p1.add_run(key)
        format_run(r1, 10.5, bold=True)
        format_paragraph(p1, after=0, line=1.05)
        p2 = cells[1].paragraphs[0]
        p2.clear()
        r2 = p2.add_run(value)
        format_run(r2, 10.5)
        format_paragraph(p2, after=0, line=1.1)
    doc.add_paragraph()


def add_cases_table(doc, title, cases):
    add_heading(doc, title, level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.autofit = False
    widths = [Inches(0.8), Inches(1.2), Inches(2.65), Inches(1.0), Inches(0.85)]
    headers = ["ID", "测试目标", "测试内容", "核心指标", "通过标准"]
    header_cells = table.rows[0].cells
    for cell, header, width in zip(header_cells, headers, widths):
        cell.width = width
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, "EAF1FB")
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(header)
        format_run(run, 10, bold=True, color="17365D")
        format_paragraph(p, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    for case in cases:
        row_cells = table.add_row().cells
        set_row_cant_split(table.rows[-1])
        for idx, width in enumerate(widths):
            row_cells[idx].width = width
            row_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(row_cells[idx], top=90, bottom=90)
        for idx, value in enumerate(case):
            p = row_cells[idx].paragraphs[0]
            p.clear()
            run = p.add_run(value)
            format_run(run, 9.5, bold=(idx == 0))
            align = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 3) else WD_ALIGN_PARAGRAPH.LEFT
            format_paragraph(p, after=0, line=1.1, align=align)
    doc.add_paragraph()


def add_footer_with_page_number(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    format_run(run, 9, color="7F7F7F")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run_page = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(r_fonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")
    r_pr.append(sz)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "7F7F7F")
    r_pr.append(color)
    run_page.append(r_pr)
    text = OxmlElement("w:t")
    text.text = "1"
    run_page.append(text)
    fld.append(run_page)
    paragraph._p.append(fld)
    run2 = paragraph.add_run(" 页")
    format_run(run2, 9, color="7F7F7F")


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    add_footer_with_page_number(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)

    add_title(
        doc,
        "美国市场 H5 测试计划（转化导向版）",
        "适用对象：占卜 / 情感 / 人生建议类 H5 产品 | 版本：V1.0 | 日期：2026-06-09",
    )

    add_body(doc, "本文档面向 H5 冷启动投放阶段，核心目标不是验证长期留存，而是尽快验证“广告点击是否能顺利转成支付”。测试重点放在点击、开始、完成、到支付页和付费率。")

    add_heading(doc, "1. 测试目标与范围")
    add_kv_table(
        doc,
        [
            ("测试目标", "验证 H5 产品在美国市场的买量转化效率，找出最优素材、人群、产品路径、价格和支付页表达。"),
            ("核心问题", "用户为什么点进来、为什么开始、为什么中途退出、为什么付费或不付费。"),
            ("测试范围", "覆盖投放侧、产品侧、用户侧、数据与合规侧四个维度，聚焦点击到支付成功的完整链路。"),
            ("阶段目标", "首轮优先看 CTR、LPV、开始率、完成率、到支付页转化率、付费率、CAC、ARPPU、退款率。"),
        ],
    )

    add_heading(doc, "2. H5 核心漏斗")
    for item in [
        "曝光 -> 点击 -> 落地页访问（LPV）",
        "LPV -> 开始测试 / 开始提问",
        "开始 -> 完成结果前流程",
        "完成 -> 到达支付页",
        "支付页 -> 支付成功",
    ]:
        add_body(doc, item, bullet=True)

    add_heading(doc, "3. 执行节奏建议")
    for text in [
        "Week 1：完成埋点、归因、H5 页面版本、素材、价格方案与支付页表达准备。",
        "Week 2：小预算并行测试平台、素材、人群和首页切入点，快速淘汰低效组合。",
        "Week 3：集中验证转化路径、价格带、结果解锁方式和支付页信任组件。",
        "Week 4：汇总各组合的付费率、CAC、ARPPU、退款率，判断是否具备继续放量条件。",
    ]:
        add_numbered(doc, text)

    add_heading(doc, "4. 判定口径")
    for text in [
        "投放层重点看 CTR、CPC、LPV 和开始率，确认广告能否把正确人群带进来。",
        "产品层重点看开始率、完成率、到支付页转化率和付费率，确认 H5 承接是否顺畅。",
        "商业层重点看 CAC、ARPPU 和退款率，确认这条链路是否具备买量价值。",
        "用户研究重点放在支付阻碍、价格敏感度和信任障碍，不把留存作为当前核心判断条件。",
    ]:
        add_body(doc, text, bullet=True)

    add_cases_table(doc, "5. 投放侧测试用例", AD_CASES)
    add_cases_table(doc, "6. 产品侧测试用例", PRODUCT_CASES)
    add_cases_table(doc, "7. 用户侧测试用例", USER_CASES)
    add_cases_table(doc, "8. 数据与合规侧测试用例", DATA_CASES)

    add_heading(doc, "9. 首轮优先测试清单")
    add_body(doc, "若预算和人力有限，建议首轮先执行以下 8 个关键用例，优先验证 H5 的成交效率：")
    for item in PRIORITY_CASES:
        add_body(doc, item, bullet=True)

    add_heading(doc, "10. 测试记录模板")
    add_kv_table(
        doc,
        [
            ("基础信息", "测试名称、用例 ID、负责人、测试时间、测试渠道 / 页面。"),
            ("实验设计", "实验组、对照组、样本量、预算、核心变量。"),
            ("关键指标", "CTR、LPV、开始率、完成率、到支付页转化率、付费率、CAC、ARPPU、退款率。"),
            ("结果结论", "是否通过、主要发现、异常说明、下一步动作。"),
        ],
    )

    add_heading(doc, "11. 结论输出要求")
    for text in [
        "哪种卖点最容易带来高质量点击，并最终转成支付。",
        "哪类人群和平台组合最适合 H5 冷启动买量。",
        "哪条页面路径最短、最顺，最能把用户推到支付页。",
        "哪种价格带和支付页表达最有机会跑出稳定付费率。",
    ]:
        add_body(doc, text, bullet=True)

    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
