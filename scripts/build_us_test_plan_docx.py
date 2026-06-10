from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = "/Users/administrator/Desktop/work/movie_work/占卜/问答平台/docs/美国市场测试计划_测试用例.docx"


AD_CASES = [
    ("AD-01", "验证主卖点", "对比 Love clarity、Self-discovery、Spiritual guidance 三种广告主题。", "CTR、CPC、LPV", "至少 1 个主题 CTR 明显领先。"),
    ("AD-02", "验证素材形式", "对比真人口播、卡牌展示、App 录屏三类素材表现。", "CTR、CVR、CPI", "找到最低 CPI 的素材形式。"),
    ("AD-03", "验证首屏钩子", "测试“Will he come back?”、“Get clarity now”、“Your personalized reading”等前 3 秒钩子。", "CTR、3 秒播放率、LPV", "至少 1 个钩子点击效率显著领先。"),
    ("AD-04", "验证平台差异", "在 Meta、TikTok、Google 三平台使用同主题素材对比。", "CPM、CTR、注册率、CAC", "筛出主力渠道和备选渠道。"),
    ("AD-05", "验证人群方向", "测试情感关系、自我成长、灵性兴趣、泛女性兴趣四类受众包。", "CTR、注册率、首单率", "确定 1 到 2 个高质量人群包。"),
    ("AD-06", "验证价格前置", "广告文案中加入“$0.99 first reading”与不带价格版本对比。", "CTR、LP 转化率、首单率", "判断价格前置是否提升流量质量。"),
    ("AD-07", "验证信任表达", "对比 Personalized guidance 与 Psychic reading 两种表达。", "审核通过率、CTR、付费率", "找到兼顾过审和转化的表达方式。"),
    ("AD-08", "验证 CTA", "测试 Start your reading、Ask your question、See your result 三种 CTA。", "CTR、注册率", "选出最强 CTA 组合。"),
]

PRODUCT_CASES = [
    ("PD-01", "验证首页定位", "首页分别强调 Tarot reading、Relationship clarity、AI guidance。", "首屏停留、开始率", "开始率最高版本胜出。"),
    ("PD-02", "验证首次体验流程", "先答题再出结果，与先输入问题再出结果两条路径对比。", "开始率、完成率、付费率", "选出完整漏斗转化最高路径。"),
    ("PD-03", "验证结果解锁方式", "直接给答案，与先给摘要再引导解锁完整结果对比。", "付费率、跳失率", "判断“先尝后买”是否更优。"),
    ("PD-04", "验证付费模式", "对比首单低价、单次付费、周订阅、月订阅。", "首付率、退款率、D7 留存", "找到转化与留存平衡点。"),
    ("PD-05", "验证价格带", "测试 $0.99、$2.99、$4.99 三档首单价格。", "首付率、ARPPU", "确认最佳价格甜点。"),
    ("PD-06", "验证信任组件", "比较加入真实评价、顾问说明、隐私承诺前后的转化变化。", "首付率、跳失率", "信任组件带来正向提升。"),
    ("PD-07", "验证支付透明度", "支付前明确展示续费说明，与弱展示版本对比。", "支付率、退款率、投诉率", "优先保留低投诉方案。"),
    ("PD-08", "验证注册节点", "先注册后体验，与游客先体验后注册两种方案对比。", "注册率、完成率、首付率", "找出阻力最小的注册策略。"),
    ("PD-09", "验证问答深度", "1 轮简版解读，与多轮追问解读对比。", "停留时长、复购率", "确认深度内容是否带动复购。"),
    ("PD-10", "验证召回机制", "测试 24 小时提醒、周运势提醒、未完成解读提醒。", "回访率、二次付费率", "选出最有效召回触发点。"),
]

USER_CASES = [
    ("US-01", "验证用户动机", "区分感情困扰、好奇尝试、长期灵性用户三类核心人群。", "首付率、留存率", "找到最高价值用户群。"),
    ("US-02", "验证需求类型", "比较用户对预测结果与情绪安慰 / 建议的偏好。", "满意度、复购率", "明确主需求方向。"),
    ("US-03", "验证内容偏好", "对比 Love、Career、Future、Self-growth 四类入口。", "点击率、完成率、付费率", "选出高需求主题。"),
    ("US-04", "验证使用时段", "观察深夜访问与白天访问用户的转化差异。", "转化率、客单价", "判断最佳投放时段。"),
    ("US-05", "验证支付障碍", "访谈已付费与未付费用户，梳理放弃支付原因。", "支付阻碍分布、首付率", "明确最大付费阻碍。"),
    ("US-06", "验证信任流失点", "定位用户在哪一步开始不信任产品。", "节点流失率、访谈反馈", "锁定最主要信任问题。"),
    ("US-07", "验证复购理由", "分析复购用户的回访原因和触发内容。", "二次购买率、回访率", "提炼复购驱动力。"),
    ("US-08", "验证退款原因", "拆解退款、差评与客服投诉文本。", "退款率、差评率", "找到首要退款原因并可优化。"),
]

DATA_CASES = [
    ("DA-01", "验证埋点完整性", "校验广告点击到支付完成全链路埋点。", "关键事件完整率", "关键事件完整率高于 95%。"),
    ("DA-02", "验证归因准确性", "核对 Meta、TikTok、Google 的转化归因口径。", "归因一致性", "误差控制在可接受范围。"),
    ("DA-03", "验证漏斗看板", "搭建注册、开始、完成、支付的日常监控看板。", "漏斗可用性", "可每日定位主要流失节点。"),
    ("CM-01", "验证广告与内容合规", "检查是否出现绝对化承诺、敏感暗示、医疗/法律/财务替代表述。", "审核通过率、投诉率", "上线素材与文案基础合规。"),
]

PRIORITY_CASES = [
    "AD-01 主卖点测试",
    "AD-02 素材形式测试",
    "AD-04 平台差异测试",
    "AD-05 人群方向测试",
    "PD-02 首次体验流程测试",
    "PD-03 结果解锁方式测试",
    "PD-04 付费模式测试",
    "PD-05 首单价格测试",
    "PD-06 信任组件测试",
    "US-01 用户分层测试",
    "US-05 支付障碍分析",
    "DA-01 埋点完整性校验",
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def format_run(run, size, bold=False, color="1F1F1F", font_name="Calibri"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def format_paragraph(paragraph, after=6, before=0, line=1.1, align=WD_ALIGN_PARAGRAPH.LEFT):
    pf = paragraph.paragraph_format
    pf.space_after = Pt(after)
    pf.space_before = Pt(before)
    pf.line_spacing = line
    paragraph.alignment = align


def add_title(doc, text, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    format_run(run, 22, bold=True, color="17365D")
    format_paragraph(p, after=8, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(subtitle)
    format_run(run2, 10, color="666666")
    format_paragraph(p2, after=18, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        format_run(run, 15, bold=True, color="2E74B5")
        format_paragraph(p, after=8, before=14)
    elif level == 2:
        format_run(run, 12.5, bold=True, color="2E74B5")
        format_paragraph(p, after=6, before=10)
    else:
        format_run(run, 11.5, bold=True, color="1F4D78")
        format_paragraph(p, after=4, before=8)


def add_body(doc, text, bullet=False):
    p = doc.add_paragraph(style="List Bullet" if bullet else None)
    run = p.add_run(text)
    format_run(run, 10.5)
    format_paragraph(p, after=6, line=1.15)


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    format_run(run, 10.5)
    format_paragraph(p, after=6, line=1.15)


def add_kv_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    widths = [Inches(1.45), Inches(5.05)]
    for key, value in rows:
        cells = table.add_row().cells
        set_row_cant_split(table.rows[-1])
        for idx, width in enumerate(widths):
            cells[idx].width = width
        set_cell_shading(cells[0], "F2F4F7")
        for cell in cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
        p1 = cells[0].paragraphs[0]
        p1.clear()
        r1 = p1.add_run(key)
        format_run(r1, 10.5, bold=True, color="1F1F1F")
        format_paragraph(p1, after=0, line=1.05)
        p2 = cells[1].paragraphs[0]
        p2.clear()
        r2 = p2.add_run(value)
        format_run(r2, 10.5)
        format_paragraph(p2, after=0, line=1.1)
    doc.add_paragraph()


def add_cases_table(doc, title, cases):
    add_heading(doc, title, level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.autofit = False
    widths = [Inches(0.8), Inches(1.15), Inches(2.7), Inches(0.95), Inches(0.9)]
    headers = ["ID", "测试目标", "测试内容", "核心指标", "通过标准"]
    header_cells = table.rows[0].cells
    for idx, (cell, header, width) in enumerate(zip(header_cells, headers, widths)):
        cell.width = width
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, "EAF1FB")
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(header)
        format_run(run, 10, bold=True, color="17365D")
        format_paragraph(p, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)

    for case in cases:
        row_cells = table.add_row().cells
        set_row_cant_split(table.rows[-1])
        for idx, width in enumerate(widths):
            row_cells[idx].width = width
            row_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(row_cells[idx], top=90, bottom=90)
        for idx, value in enumerate(case):
            p = row_cells[idx].paragraphs[0]
            p.clear()
            run = p.add_run(value)
            format_run(run, 9.5, bold=(idx == 0))
            align = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 3) else WD_ALIGN_PARAGRAPH.LEFT
            format_paragraph(p, after=0, line=1.1, align=align)
    doc.add_paragraph()


def add_footer_with_page_number(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    format_run(run, 9, color="7F7F7F")

    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run_page = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(r_fonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")
    r_pr.append(sz)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "7F7F7F")
    r_pr.append(color)
    run_page.append(r_pr)
    text = OxmlElement("w:t")
    text.text = "1"
    run_page.append(text)
    fld.append(run_page)
    paragraph._p.append(fld)

    run2 = paragraph.add_run(" 页")
    format_run(run2, 9, color="7F7F7F")


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)
    add_footer_with_page_number(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)

    add_title(
        doc,
        "美国市场测试计划（测试用例版）",
        "适用对象：占卜 / 情感 / 人生建议类问答产品 | 版本：V1.0 | 日期：2026-06-09",
    )

    add_body(doc, "本文档用于支持美国市场冷启动测试，重点验证获客效率、站内转化、用户价值和合规可行性。建议以 3 到 4 周为首轮验证周期，以“能否稳定买量并形成正向复购信号”为核心判断标准。")

    add_heading(doc, "1. 测试目标与范围", level=1)
    add_kv_table(
        doc,
        [
            ("测试目标", "验证这款占卜/问答产品在美国市场是否具备规模化投放基础，并明确最优人群、素材、产品路径和付费方式。"),
            ("测试范围", "覆盖投放侧、产品侧、用户侧、数据与合规侧四个维度，兼顾获客、转化、留存与风控。"),
            ("测试周期", "建议首轮执行 3 到 4 周：第 1 周准备，第 2 周探索，第 3 周集中验证，第 4 周复盘与放量判断。"),
            ("阶段目标", "首轮不以 ROAS 为唯一目标，优先看 CTR、注册率、首次问答完成率、首单率、D7 留存、退款率。"),
        ],
    )

    add_heading(doc, "2. 执行节奏建议", level=1)
    for text in [
        "Week 1：完成埋点、归因、看板、广告素材、落地页和首单价格方案准备，并完成基础合规审核。",
        "Week 2：小预算并行测试素材、人群、平台和落地页，快速淘汰表现靠后的组合。",
        "Week 3：将预算集中到前 20% 组合，同时测试首屏路径、付费模式与信任组件。",
        "Week 4：评估 D7 留存、二次付费率、退款率和渠道放量潜力，给出 go / no-go 判断。",
    ]:
        add_numbered(doc, text)

    add_heading(doc, "3. 判定口径", level=1)
    for text in [
        "投放层重点看 CTR、CPC、LPV、注册率和 CAC，避免只看点击不看站内质量。",
        "产品层重点看首次核心行为完成率、首单率、退款率和信任流失点。",
        "用户层重点看 D1 / D7 留存、二次付费率、满意度和主要退款原因。",
        "最终保留的方案需要同时满足“能点进来、能留下来、愿意付费、风险可控”。",
    ]:
        add_body(doc, text, bullet=True)

    add_cases_table(doc, "4. 投放侧测试用例", AD_CASES)
    add_cases_table(doc, "5. 产品侧测试用例", PRODUCT_CASES)
    add_cases_table(doc, "6. 用户侧测试用例", USER_CASES)
    add_cases_table(doc, "7. 数据与合规侧测试用例", DATA_CASES)

    add_heading(doc, "8. 首轮优先测试清单", level=1)
    add_body(doc, "若资源有限，建议优先执行以下 12 个用例，以最快速度判断美国市场是否值得继续放大：")
    for item in PRIORITY_CASES:
        add_body(doc, item, bullet=True)

    add_heading(doc, "9. 测试记录模板", level=1)
    add_body(doc, "建议每个实验统一按下列字段沉淀，方便后续做横向复盘和渠道归因：")
    add_kv_table(
        doc,
        [
            ("基础信息", "测试名称、用例 ID、负责人、测试时间、测试渠道 / 页面。"),
            ("实验设计", "实验组、对照组、样本量、预算、核心变量。"),
            ("观察指标", "CTR、注册率、完成率、首单率、CAC、D7 留存、退款率等。"),
            ("结果结论", "是否通过、主要发现、异常说明、下一步动作。"),
        ],
    )

    add_heading(doc, "10. 结论输出要求", level=1)
    for text in [
        "美国用户最吃哪种切入角度：情感、灵性、自我探索还是陪伴建议。",
        "哪一类渠道最适合作为主力买量渠道，哪一类渠道更适合补量或品牌曝光。",
        "用户真正买单的是预测结果本身，还是被理解、被安慰和被引导的过程。",
        "该产品更适合低价冲量模式，还是高信任、慢转化、靠复购拉长生命周期的模式。",
    ]:
        add_body(doc, text, bullet=True)

    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
